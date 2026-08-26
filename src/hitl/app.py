"""
FinAlze HITL Validation UI

Streamlit-based human-in-the-loop interface for:
- Viewing extracted data alongside source documents
- Editing and correcting OCR results
- Validating with accounting rules
- Approving/rejecting extractions
"""

import streamlit as st
import pandas as pd
import json
from pathlib import Path
from datetime import datetime
import sys
import io
import time

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from pipeline import FinAlzePipeline, PipelineConfig, ExtractionResult
from accounting.statement_segmenter import segment_tables
from accounting.validator import ValidationResult, ValidationReport

# Display name mappings for statement types
_STATEMENT_DISPLAY = {
    "bilan": "Bilan",
    "compte_resultat": "Compte de Resultat",
    "flux_tresorerie": "Flux de Tresorerie",
}
_SUB_TYPE_DISPLAY = {"actif": "Actif", "passif": "Passif"}


def _reconstruct_report(report_dict: dict) -> ValidationReport:
    """Reconstruct a ValidationReport from the dict stored in metadata."""
    report = ValidationReport()
    for check_data in report_dict.get("checks", []):
        vr = ValidationResult(
            check_name=check_data["name"],
            passed=check_data["passed"],
            severity=check_data.get("severity", "WARNING"),
            message=check_data.get("message", ""),
            expected=check_data.get("expected"),
            actual=check_data.get("actual"),
            difference=check_data.get("difference"),
        )
        report.add_check(vr)
    report.compute_severity_weighted_score()
    return report


# Page configuration
st.set_page_config(
    page_title="FinAlze OCR Validator",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: 700;
        color: #1f77b4;
        margin-bottom: 1rem;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        border: 1px solid #ffeeba;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .error-box {
        background-color: #f8d7da;
        border: 1px solid #f5c6cb;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)


def init_session_state():
    """Initialize session state variables."""
    if 'extraction_result' not in st.session_state:
        st.session_state.extraction_result = None
    if 'edited_tables' not in st.session_state:
        st.session_state.edited_tables = []
    if 'validation_history' not in st.session_state:
        st.session_state.validation_history = []
    if 'pipeline' not in st.session_state:
        st.session_state.pipeline = None
    if 'current_file' not in st.session_state:
        st.session_state.current_file = None
    if 'last_processing_time' not in st.session_state:
        st.session_state.last_processing_time = None


def render_sidebar():
    """Render the sidebar with configuration options."""
    st.sidebar.markdown("## ⚙️ Configuration")
    
    # QCS Threshold sliders
    st.sidebar.markdown("### QCS Thresholds")
    tau_fast = st.sidebar.slider(
        "Tier 1 (Fast OCR)",
        min_value=0.5,
        max_value=1.0,
        value=0.75,
        step=0.05,
        help="Minimum QCS for Tier 1 to accept without escalation"
    )
    tau_vlm = st.sidebar.slider(
        "Tier 2 (VLM)",
        min_value=0.5,
        max_value=1.0,
        value=0.70,
        step=0.05,
        help="Minimum QCS for Tier 2 to accept without HITL"
    )
    
    # Preprocessing options
    st.sidebar.markdown("### Preprocessing")
    preprocess_resize = st.sidebar.toggle(
        "Resize Large Images",
        value=True,
        help="Resize images wider than 2000px for faster processing"
    )
    preprocess_grayscale = st.sidebar.toggle(
        "Convert to Grayscale",
        value=True,
        help="Convert to grayscale for faster OCR processing"
    )
    max_image_width = st.sidebar.number_input(
        "Max Image Width (px)",
        min_value=500,
        max_value=4000,
        value=2000,
        step=100,
        help="Maximum width before resizing"
    )
    
    # Accounting standard
    st.sidebar.markdown("### Accounting Standard")
    standard = st.sidebar.selectbox(
        "Select Standard",
        options=["IFRS", "NCT", "SYSCOHADA"],
        index=0
    )
    
    # Auto-apply: always keep pipeline in sync with current slider values
    current_config_key = f"{tau_fast}_{tau_vlm}_{standard}"
    if 'last_config_key' not in st.session_state:
        st.session_state.last_config_key = None
    
    # Recreate pipeline if config changed or pipeline doesn't exist
    if st.session_state.pipeline is None or st.session_state.last_config_key != current_config_key:
        config = PipelineConfig(
            tau_fast_high=tau_fast,
            tau_vlm_med=tau_vlm,
            accounting_standard=standard
        )
        st.session_state.pipeline = FinAlzePipeline(config)
        st.session_state.last_config_key = current_config_key
        st.sidebar.caption(f"✓ Thresholds: T1={tau_fast:.2f}, T2={tau_vlm:.2f}")
    
    return tau_fast, tau_vlm, standard, preprocess_resize, preprocess_grayscale, max_image_width


def render_upload_section():
    """Render the file upload section."""
    st.markdown("### 📁 Upload Document")
    
    uploaded_file = st.file_uploader(
        "Upload a financial statement (PDF, DOCX, JPG, PNG)",
        type=["pdf", "docx", "jpg", "jpeg", "png"],
        help="Supported formats: PDF, DOCX, JPG, PNG"
    )
    
    # Or select from Samples directory
    st.markdown("**Or select from samples:**")
    samples_dir = Path(__file__).parent.parent.parent / "Samples"
    if samples_dir.exists():
        sample_files = list(samples_dir.glob("*"))
        sample_names = ["-- Select --"] + [f.name for f in sample_files if f.is_file()]
        selected_sample = st.selectbox("Sample files", sample_names)
        
        if selected_sample != "-- Select --":
            return samples_dir / selected_sample
    
    return uploaded_file


def process_document(file_input, preprocess_resize=True, preprocess_grayscale=True, max_image_width=2000, accounting_standard="NCT"):
    """Process the uploaded document through the pipeline."""
    if st.session_state.pipeline is None:
        st.session_state.pipeline = FinAlzePipeline()
    
    # Handle both file objects and paths
    if isinstance(file_input, Path):
        input_path = file_input
    else:
        # Save uploaded file temporarily
        temp_dir = Path(__file__).parent.parent.parent / "output" / "temp"
        temp_dir.mkdir(parents=True, exist_ok=True)
        temp_path = temp_dir / file_input.name
        temp_path.write_bytes(file_input.getvalue())
        input_path = temp_path
    
    st.session_state.current_file = input_path
    
    with st.spinner(f"Processing {input_path.name} ({accounting_standard})..."):
        start_time = time.perf_counter()
        import os
        api_url = os.environ.get("FINALZE_API_URL")
        if api_url:
            # API mode: submit to the backend and reconstruct an ExtractionResult from the JSON.
            from client.sdk import FinAlzeClient
            with FinAlzeClient(base_url=api_url, api_key=os.environ.get("FINALZE_API_KEY"),
                               token=os.environ.get("FINALZE_API_TOKEN")) as _client:
                d = _client.submit_and_wait(input_path, accounting_standard=accounting_standard)
            result = ExtractionResult(
                text=d.get("text", "") or "",
                tables=d.get("tables") or [],
                qcs_score=d.get("qcs_score") or 0.0,
                tier_used=d.get("tier_used") or 0,
                confidence_details=d.get("confidence_details") or {},
                metadata=d.get("metadata") or {},
            )
        else:
            # Local mode (default): run the pipeline in-process.
            result = st.session_state.pipeline.process_document(
                input_path,
                preprocess_resize=preprocess_resize,
                preprocess_grayscale=preprocess_grayscale,
                max_image_width=max_image_width,
                accounting_standard=accounting_standard
            )
        elapsed_time = time.perf_counter() - start_time
        st.session_state.last_processing_time = elapsed_time
        st.session_state.extraction_result = result
        
        # Convert tables to DataFrames for editing, with auto-correction applied
        st.session_state.edited_tables = []
        try:
            from accounting.validator import auto_correct_variation_percentages
            apply_correction = True
        except ImportError:
            apply_correction = False
        
        for table_data in result.tables:
            if isinstance(table_data, list):
                df = pd.DataFrame(table_data)
            else:
                df = pd.DataFrame(table_data)
            
            # Apply auto-correction immediately so UI shows corrected values
            if apply_correction:
                df, _ = auto_correct_variation_percentages(df)
            
            st.session_state.edited_tables.append(df)
    
    return result, elapsed_time


def render_source_view():
    """Render the source document view."""
    st.markdown("### 📄 Source Document")
    
    if st.session_state.current_file:
        file_path = st.session_state.current_file
        suffix = file_path.suffix.lower()
        
        if suffix in ['.jpg', '.jpeg', '.png']:
            st.image(str(file_path), caption=file_path.name, use_container_width=True)
        elif suffix == '.pdf':
            try:
                import fitz  # PyMuPDF
                pdf_doc = fitz.open(str(file_path))
                num_pages = len(pdf_doc)
                if num_pages > 1:
                    page_num = st.selectbox(
                        "Page",
                        options=list(range(1, num_pages + 1)),
                        format_func=lambda x: f"Page {x} of {num_pages}",
                        key="pdf_page_selector",
                    )
                else:
                    page_num = 1
                page = pdf_doc[page_num - 1]
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("png")
                st.image(img_bytes, caption=f"{file_path.name} — Page {page_num}/{num_pages}", use_container_width=True)
                pdf_doc.close()
            except ImportError:
                st.warning("PyMuPDF not installed. Install with: `pip install PyMuPDF`")
                st.markdown(f"**File:** `{file_path}`")
            except Exception as e:
                st.error(f"Could not render PDF: {e}")
                st.markdown(f"**File:** `{file_path}`")
        elif suffix == '.docx':
            st.info("DOCX preview not available. Showing extracted images.")
            # Try to show embedded images
            try:
                from docx import Document
                doc = Document(file_path)
                for i, rel in enumerate(doc.part.rels.values()):
                    if "image" in rel.reltype:
                        image_data = rel.target_part.blob
                        st.image(image_data, caption=f"Embedded Image {i+1}", use_container_width=True)
            except Exception as e:
                st.error(f"Could not load DOCX: {e}")


def render_extraction_view():
    """Render the extraction results with editable tables."""
    st.markdown("### 📊 Extracted Data")
    
    result = st.session_state.extraction_result
    if result is None:
        st.info("No extraction result yet. Upload a document to start.")
        return
    
    # Metrics row
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        # Show QCS score with grade badge
        qcs_grade = result.confidence_details.get('qcs_grade', 'unknown')
        grade_colors = {
            'excellent': '🟢',
            'good': '🟡',
            'fair': '🟠',
            'poor': '🔴',
            'unknown': '⚪',
        }
        grade_icon = grade_colors.get(qcs_grade, '⚪')
        st.metric("QCS Score", f"{result.qcs_score:.2f} {grade_icon}", 
                  help=f"Grade: {qcs_grade.upper()}")
    with col2:
        st.metric("Tier Used", f"Tier {result.tier_used}")
    with col3:
        st.metric("Tables Found", len(result.tables))
    with col4:
        hitl_required = result.metadata.get('hitl_required', False)
        hitl_reason = result.confidence_details.get('hitl_reason', '')
        st.metric("HITL Required", "Yes" if hitl_required else "No")
        if hitl_required and hitl_reason:
            st.caption(f"Reason: {hitl_reason}")
    
    # Tier Attempts Display
    tier_attempts = result.confidence_details.get('tier_attempts', [])
    if tier_attempts:
        with st.expander("Tier QCS Breakdown", expanded=True):
            for attempt in tier_attempts:
                tier_num = attempt.get('tier', '?')
                tier_name = attempt.get('name', 'Unknown')
                qcs = attempt.get('qcs', 0)
                passed = attempt.get('passed', False)
                
                # Color-coded status
                if passed:
                    status = "✅ Passed"
                    color = "green"
                else:
                    status = "❌ Below threshold"
                    color = "red"
                
                col_tier, col_qcs, col_status = st.columns([2, 1, 2])
                with col_tier:
                    st.markdown(f"**Tier {tier_num}:** {tier_name}")
                with col_qcs:
                    st.markdown(f"QCS: **{qcs:.3f}**")
                with col_status:
                    st.markdown(f":{color}[{status}]")
    
    # OCR Correction Stats
    correction_stats = result.metadata.get('correction_stats', {})
    total_corr = correction_stats.get('total_corrections', 0)
    if total_corr > 0:
        with st.expander(f"OCR Corrections Applied ({total_corr})"):
            stat_items = [
                ('term_corrections', 'Term fixes'),
                ('camelcase_fixes', 'CamelCase splits'),
                ('missing_space_fixes', 'Space fixes'),
                ('ocr_char_fixes', 'Char substitutions'),
                ('underline_fixes', 'Underline fixes'),
            ]
            active = [(k, d) for k, d in stat_items if correction_stats.get(k, 0) > 0]
            if active:
                cols = st.columns(len(active))
                for j, (key, display) in enumerate(active):
                    cols[j].metric(display, correction_stats[key])

    # Confidence details (collapsed)
    with st.expander("Full Confidence Details"):
        st.json(result.confidence_details)

    # Editable tables — labelled by statement classification
    st.markdown("#### Edit Extracted Tables")

    # Segment tables for classification labels (cached in session state)
    if 'table_segments' not in st.session_state or len(st.session_state.get('table_segments', [])) != len(st.session_state.edited_tables):
        try:
            st.session_state.table_segments = segment_tables(st.session_state.edited_tables)
        except Exception:
            st.session_state.table_segments = []

    for i, df in enumerate(st.session_state.edited_tables):
        # Build label from segment classification
        label = f"Table {i+1}"
        seg = st.session_state.table_segments[i] if i < len(st.session_state.table_segments) else None
        if seg and seg.statement_type in _STATEMENT_DISPLAY:
            label += f" — {_STATEMENT_DISPLAY[seg.statement_type]}"
            if seg.sub_type in _SUB_TYPE_DISPLAY:
                label += f" ({_SUB_TYPE_DISPLAY[seg.sub_type]})"

        col_label, col_badge = st.columns([5, 1])
        with col_label:
            st.markdown(f"**{label}**")
        with col_badge:
            if seg and seg.statement_type != "unknown":
                badge_color = "green" if seg.confidence >= 0.6 else "orange"
                st.badge(f"{seg.confidence:.0%}", color=badge_color)
        
        edited_df = st.data_editor(
            df,
            num_rows="dynamic",
            use_container_width=True,
            key=f"table_{i}"
        )
        st.session_state.edited_tables[i] = edited_df
    
    # Raw text view
    with st.expander("Raw Extracted Text"):
        st.text_area("Markdown Output", result.text, height=300)


def render_validation_section():
    """Render the validation and approval section with automatic checks."""
    st.markdown("### ✅ Validation")
    
    result = st.session_state.extraction_result
    if result is None:
        return
    
    # ---------------------------------------------------------------
    # Build validation report: prefer pipeline's V1-V4, fall back
    # ---------------------------------------------------------------
    try:
        from accounting.validator import validate_extraction, auto_correct_variation_percentages

        tables = st.session_state.edited_tables

        # Auto-correct variation percentages
        all_corrections = []
        corrected_tables = []
        for i, df in enumerate(tables):
            df_corrected, corrections = auto_correct_variation_percentages(df)
            corrected_tables.append(df_corrected)
            for c in corrections:
                c['table_idx'] = i
                all_corrections.append(c)

        if all_corrections:
            st.session_state.edited_tables = corrected_tables
            tables = corrected_tables
            with st.expander(f"Auto-corrected {len(all_corrections)} variation %", expanded=False):
                for c in all_corrections:
                    st.markdown(f"- **{c['label']}**: `{c['original']}` -> `{c['corrected']}`")

        # Try pipeline's pre-computed V1-V4 report, then run V1-V4 live, then basic
        pipeline_report_dict = result.metadata.get('validation_report')
        if pipeline_report_dict and isinstance(pipeline_report_dict, dict):
            validation_report = _reconstruct_report(pipeline_report_dict)
        else:
            try:
                from accounting.rule_engine import validate_with_rules
                standard = result.metadata.get('accounting_standard', 'NCT')
                validation_report = validate_with_rules(tables, standard=standard)
            except ImportError:
                validation_report = validate_extraction(tables)

        st.session_state.validation_report = validation_report

    except ImportError:
        validation_report = None
        st.warning("Automatic validation module not available")

    # ---------------------------------------------------------------
    # Validation summary bar
    # ---------------------------------------------------------------
    if validation_report:
        sws = validation_report.severity_weighted_score
        crit = validation_report.critical_failures
        summary_cols = st.columns([1, 1, 1, 2])
        with summary_cols[0]:
            st.metric("Checks Passed", f"{validation_report.passed_checks}/{validation_report.total_checks}")
        with summary_cols[1]:
            sws_display = f"{sws:.0%}" if sws is not None else "N/A"
            st.metric("Severity Score", sws_display, help="Weighted: CRITICAL=3x, ERROR=2x, WARNING=1x")
        with summary_cols[2]:
            st.metric("Critical Failures", crit)
            if crit > 0:
                st.caption(":red[Forces HITL review]")
        with summary_cols[3]:
            if validation_report.overall_passed:
                st.success("All checks passed")
            elif crit > 0:
                st.error(f"{crit} critical failure(s) detected")
            else:
                st.warning(f"{validation_report.failed_checks} check(s) failed")

    # ---------------------------------------------------------------
    # Tier-grouped checks + Quality assessment side by side
    # ---------------------------------------------------------------
    col1, col2 = st.columns(2)

    with col1:
        if 'check_overrides' not in st.session_state:
            st.session_state.check_overrides = {}

        if validation_report:
            # Determine what statement types were found (for "not applicable" messages)
            segments = st.session_state.get('table_segments', [])
            found_types = {s.statement_type for s in segments} if segments else set()
            has_multi_statements = len(found_types - {"unknown"}) > 1
            has_bilan = "bilan" in found_types

            # Group checks by validation tier
            tier_meta = {
                "General": {
                    "checks": [],
                    "desc": "Basic extraction checks",
                    "empty_reason": None,
                },
                "V1 — Intra-Statement": {
                    "checks": [],
                    "desc": "Balance equations, subtotals, VNC, CR equation",
                    "empty_reason": None,
                },
                "V2 — Cross-Statement": {
                    "checks": [],
                    "desc": "Bilan <-> CR <-> TFT coherence",
                    "empty_reason": (
                        "Requires multiple statement types (Bilan + CR or TFT). "
                        f"Only found: {', '.join(found_types - {'unknown'}) or 'unknown'}."
                    ) if not has_multi_statements else None,
                },
                "V3 — Anomaly Detection": {
                    "checks": [],
                    "desc": "Negative values, null CA, magnitude mismatches",
                    "empty_reason": "No anomalies detected." if has_bilan else "No bilan or CR tables found.",
                },
                "V4 — Ratios": {
                    "checks": [],
                    "desc": "Fonds de Roulement, Tresorerie Nette",
                    "empty_reason": (
                        "Requires Bilan with identifiable totals (Actifs NC, Capitaux Propres, Passifs NC)."
                    ) if not has_bilan else "Could not locate required totals in Bilan.",
                },
            }
            for check in validation_report.checks:
                name = check.check_name
                if name.startswith("V2:") or ("Match" in name and any(kw in name for kw in ["Bilan", "CR", "TFT"])):
                    tier_meta["V2 — Cross-Statement"]["checks"].append(check)
                elif name.startswith("V3:") or any(kw in name for kw in ["Negative", "Null CA", "Magnitude"]):
                    tier_meta["V3 — Anomaly Detection"]["checks"].append(check)
                elif name.startswith("V4:") or "Fonds de Roulement" in name:
                    tier_meta["V4 — Ratios"]["checks"].append(check)
                elif any(kw in name for kw in ["Balance", "VNC", "CR Equation", "Row Sum", "Var%", "Field Type"]):
                    tier_meta["V1 — Intra-Statement"]["checks"].append(check)
                else:
                    tier_meta["General"]["checks"].append(check)

            _SEV_COLOR = {"CRITICAL": "red", "ERROR": "orange", "WARNING": "blue"}
            check_idx = 0

            for tier_name, meta in tier_meta.items():
                checks = meta["checks"]

                if checks:
                    tier_passed = sum(1 for c in checks if c.passed)
                    tier_total = len(checks)
                    tier_ok = tier_passed == tier_total

                    with st.expander(f"{tier_name}  ({tier_passed}/{tier_total})", expanded=not tier_ok):
                        st.caption(meta["desc"])
                        for check in checks:
                            check_key = f"check_{check_idx}_{check.check_name}"
                            default_value = st.session_state.check_overrides.get(check_key, check.passed)

                            icon = ":green[PASS]" if default_value else ":red[FAIL]"
                            sev_color = _SEV_COLOR.get(check.severity, "gray")

                            c1, c2, c3 = st.columns([0.5, 4, 1])
                            with c1:
                                st.markdown(icon)
                            with c2:
                                new_value = st.checkbox(
                                    check.check_name,
                                    value=default_value,
                                    key=check_key,
                                    help=check.message or None,
                                    label_visibility="visible",
                                )
                            with c3:
                                st.badge(check.severity, color=sev_color)

                            if new_value != check.passed:
                                st.session_state.check_overrides[check_key] = new_value
                                st.caption(f"  *Overridden* (auto: {'pass' if check.passed else 'fail'})")
                            elif check.message and not check.passed:
                                st.caption(f"  {check.message}")

                            check_idx += 1
                else:
                    # Always show all tiers — explain why empty
                    if tier_name == "General":
                        continue  # General with no checks is fine to hide
                    with st.expander(f"{tier_name}  (--)", expanded=False):
                        st.caption(meta["desc"])
                        reason = meta.get("empty_reason") or "No checks produced."
                        st.info(reason)

        # Manual verification
        st.markdown("---")
        st.markdown("#### Manual Verification")
        st.caption("Requires human review")
        manual_visual = st.checkbox("Visual accuracy verified", value=False, key="manual_visual",
                                     help="Confirm extracted data matches the source document visually")
        manual_values = st.checkbox("All values match source", value=False, key="manual_values",
                                     help="Confirm all numeric values were extracted correctly")
        manual_rows = st.checkbox("No missing rows", value=False, key="manual_rows",
                                   help="Confirm no table rows were skipped or merged incorrectly")

    with col2:
        st.markdown("#### Quality Assessment")

        # Auto-suggest quality based on validation results
        if validation_report:
            if validation_report.overall_passed and result.qcs_score >= 0.85:
                default_quality = 0
            elif validation_report.overall_passed and result.qcs_score >= 0.75:
                default_quality = 1
            elif validation_report.overall_passed:
                default_quality = 2
            else:
                default_quality = 3
        else:
            default_quality = 1

        quality = st.radio(
            "Overall Quality",
            options=["Excellent", "Good", "Acceptable", "Needs Correction", "Reject"],
            index=default_quality
        )

        notes = st.text_area("Validation Notes", placeholder="Add any notes about this extraction...")
    
    # Action buttons
    st.markdown("---")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("✅ Approve & Save", type="primary", use_container_width=True):
            save_validation("approved", quality, notes)
            st.success("Extraction approved and saved!")
    
    with col2:
        if st.button("🔄 Re-process", use_container_width=True):
            if st.session_state.current_file:
                process_document(st.session_state.current_file)
                st.rerun()
    
    with col3:
        if st.button("❌ Reject", use_container_width=True):
            save_validation("rejected", quality, notes)
            st.warning("Extraction rejected.")


def save_validation(status: str, quality: str, notes: str):
    """Save the validation result."""
    result = st.session_state.extraction_result
    
    # Build the full validation report dict (with user overrides applied)
    saved_report = None
    if st.session_state.get('validation_report'):
        vr = st.session_state.validation_report
        overrides = st.session_state.get('check_overrides', {})
        saved_report = {
            "overall_passed": vr.overall_passed,
            "total_checks": vr.total_checks,
            "passed_checks": vr.passed_checks,
            "failed_checks": vr.failed_checks,
            "critical_failures": vr.critical_failures,
            "severity_weighted_score": (
                round(vr.severity_weighted_score, 3)
                if vr.severity_weighted_score is not None else None
            ),
            "checks": [
                {
                    "name": c.check_name,
                    "passed": overrides.get(f"check_{i}_{c.check_name}", c.passed),
                    "auto_passed": c.passed,
                    "overridden": f"check_{i}_{c.check_name}" in overrides,
                    "severity": c.severity,
                    "message": c.message,
                    "expected": c.expected,
                    "actual": c.actual,
                    "difference": c.difference,
                }
                for i, c in enumerate(vr.checks)
            ],
        }

    # Build table segment info
    segment_info = []
    for seg in st.session_state.get('table_segments', []):
        segment_info.append({
            "table_index": seg.table_index,
            "statement_type": seg.statement_type,
            "sub_type": seg.sub_type,
            "confidence": round(seg.confidence, 3),
        })

    validation_record = {
        "timestamp": datetime.now().isoformat(),
        "file": str(st.session_state.current_file) if st.session_state.current_file else None,
        "status": status,
        "quality": quality,
        "notes": notes,
        "qcs_score": result.qcs_score if result else None,
        "tier_used": result.tier_used if result else None,
        "confidence_details": result.confidence_details if result else None,
        "correction_stats": result.metadata.get('correction_stats') if result else None,
        "validation_report": saved_report,
        "table_segments": segment_info,
        "tables": [df.to_dict(orient='records') for df in st.session_state.edited_tables],
    }
    
    st.session_state.validation_history.append(validation_record)
    
    # Save to file
    output_dir = Path(__file__).parent.parent.parent / "output" / "validations"
    output_dir.mkdir(parents=True, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"validation_{timestamp}.json"
    
    with open(output_dir / filename, 'w', encoding='utf-8') as f:
        json.dump(validation_record, f, indent=2, ensure_ascii=False)


def render_history_section():
    """Render the validation history section."""
    st.markdown("### 📜 Validation History")
    
    if not st.session_state.validation_history:
        st.info("No validations yet this session.")
        return
    
    for i, record in enumerate(reversed(st.session_state.validation_history)):
        status_icon = "✅" if record["status"] == "approved" else "❌"
        with st.expander(f"{status_icon} {record.get('file', 'Unknown')} - {record['timestamp'][:19]}"):
            st.json(record)


def main():
    """Main application entry point."""
    init_session_state()
    
    # Header
    st.markdown('<p class="main-header">📊 FinAlze OCR Validator</p>', unsafe_allow_html=True)
    st.markdown("Human-in-the-loop validation for financial statement extraction")
    
    # Sidebar
    tau_fast, tau_vlm, standard, preprocess_resize, preprocess_grayscale, max_image_width = render_sidebar()
    
    # Main content - simplified to 2 tabs
    tab1, tab2 = st.tabs(["🔍 Validate", "📜 History"])
    
    with tab1:
        # Upload section at top
        st.markdown("### 📁 Select Document")
        
        upload_col1, upload_col2 = st.columns([2, 1])
        
        with upload_col1:
            uploaded_file = st.file_uploader(
                "Upload a financial statement",
                type=["pdf", "docx", "jpg", "jpeg", "png"],
                help="Supported: PDF, DOCX, JPG, PNG",
                label_visibility="collapsed"
            )
            
            # Sample selector
            samples_dir = Path(__file__).parent.parent.parent / "Samples"
            file_input = None
            if samples_dir.exists():
                sample_files = list(samples_dir.glob("*"))
                sample_names = ["-- Or select sample --"] + [f.name for f in sample_files if f.is_file()]
                selected_sample = st.selectbox("Samples", sample_names, label_visibility="collapsed")
                
                if selected_sample != "-- Or select sample --":
                    file_input = samples_dir / selected_sample
            
            if uploaded_file:
                file_input = uploaded_file
        
        with upload_col2:
            if file_input:
                if st.button("🚀 Process", type="primary", use_container_width=True):
                    result, elapsed = process_document(
                        file_input,
                        preprocess_resize=preprocess_resize,
                        preprocess_grayscale=preprocess_grayscale,
                        max_image_width=max_image_width,
                        accounting_standard=standard
                    )
                    st.success(f"Done in {elapsed:.1f}s! QCS: {result.qcs_score:.2f} ({standard})")
        
        st.markdown("---")
        
        # Side-by-side: Source Document | Extracted Data
        source_col, extract_col = st.columns([1, 1])
        
        with source_col:
            render_source_view()
        
        with extract_col:
            render_extraction_view()
        
        # Validation section below
        if st.session_state.extraction_result:
            st.markdown("---")
            render_validation_section()
    
    with tab2:
        render_history_section()


if __name__ == "__main__":
    main()
