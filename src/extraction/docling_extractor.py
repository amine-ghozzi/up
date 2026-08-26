"""
Docling-based table extraction for financial documents.

Uses DoclingTableFormer for accurate table structure recognition
with EasyOCR for French text recognition.
"""

import logging
from pathlib import Path
from typing import List, Optional, Tuple
from dataclasses import dataclass, field
import io

from PIL import Image
import pandas as pd

logger = logging.getLogger(__name__)


@dataclass
class TableExtractionResult:
    """Result from table extraction."""
    tables: List[pd.DataFrame] = field(default_factory=list)
    text: str = ""
    confidence: float = 0.0
    layout_score: float = 0.0
    ocr_score: float = 0.0
    corrections_made: int = 0  # For QCS penalty calculation
    # Grade-mapped fields (Docling grades recommended over raw scores)
    ocr_grade: str = "FAIR"
    layout_grade: str = "FAIR"
    parse_grade: str = "FAIR"
    low_grade: str = "FAIR"  # Worst page grade — POOR triggers HITL pre-gate
    metadata: dict = field(default_factory=dict)


class DoclingExtractor:
    """
    Docling-based document extractor with TableFormer.
    
    Optimized for French financial documents:
    - Balance sheets (Bilan)
    - Income statements (Compte de Résultat)
    - Cash flow statements (Flux de Trésorerie)
    """
    
    # Class-level singleton for converter (shared across all instances)
    _shared_converter = None
    _shared_converter_threshold = None  # Track threshold used for singleton

    def __init__(
        self,
        language: str = "fr",
        confidence_threshold: float = 0.5,
        use_accurate_mode: bool = True,
        preprocess_resize: bool = True,
        preprocess_grayscale: bool = True,
        max_image_width: int = 2000,
        accounting_standard: str = "NCT"
    ):
        self.language = language
        self.confidence_threshold = confidence_threshold
        self.use_accurate_mode = use_accurate_mode
        self.preprocess_resize = preprocess_resize
        self.preprocess_grayscale = preprocess_grayscale
        self.max_image_width = max_image_width
        self.accounting_standard = accounting_standard

    @classmethod
    def _get_converter(cls, confidence_threshold: float = 0.5):
        """Lazy-load the Docling converter (class-level singleton).

        Passes pipeline options including OCR confidence threshold.
        Reinitializes if threshold changes.
        """
        if cls._shared_converter is not None and cls._shared_converter_threshold == confidence_threshold:
            return cls._shared_converter

        logger.info("Initializing Docling converter (first run may download models)...")

        from docling.document_converter import DocumentConverter

        # Use default converter — already supports PDF, IMAGE, DOCX
        # with OCR and table structure enabled out of the box.
        # Custom PdfPipelineOptions / FormatOption wiring is version-sensitive,
        # so we rely on Docling defaults which work well for financial docs.
        cls._shared_converter = DocumentConverter()
        cls._shared_converter_threshold = confidence_threshold
        logger.info("Docling converter initialized (singleton)")

        return cls._shared_converter
    
    def _preprocess_image(self, image: Image.Image) -> Image.Image:
        """
        Preprocess image for faster OCR.
        
        Args:
            image: PIL Image
            
        Returns:
            Preprocessed PIL Image
        """
        original_size = image.size
        
        # Resize if image is too large
        if self.preprocess_resize and image.width > self.max_image_width:
            ratio = self.max_image_width / image.width
            new_height = int(image.height * ratio)
            image = image.resize((self.max_image_width, new_height), Image.Resampling.LANCZOS)
            logger.info(f"Resized image from {original_size} to {image.size}")
        
        # Convert to grayscale for faster processing
        if self.preprocess_grayscale and image.mode != 'L':
            image = image.convert('L')
            logger.info("Converted image to grayscale")
        
        return image
    
    def extract_from_pdf(self, pdf_path: Path) -> TableExtractionResult:
        """
        Extract tables and text from a PDF.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            TableExtractionResult with extracted data
        """
        converter = self._get_converter(self.confidence_threshold)

        logger.info(f"Extracting from PDF: {pdf_path}")
        result = converter.convert(str(pdf_path))
        
        return self._process_result(result)
    
    def extract_from_image(self, image: Image.Image) -> TableExtractionResult:
        """
        Extract tables and text from an image.
        
        Args:
            image: PIL Image
            
        Returns:
            TableExtractionResult with extracted data
        """
        converter = self._get_converter(self.confidence_threshold)

        # Preprocess image for faster OCR
        image = self._preprocess_image(image)
        
        # Save image temporarily
        import tempfile
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
            temp_path = Path(f.name)
            image.save(temp_path)
        
        try:
            result = converter.convert(str(temp_path))
            return self._process_result(result)
        finally:
            temp_path.unlink()
    
    def extract_from_docx(self, docx_path: Path) -> TableExtractionResult:
        """
        Extract tables from DOCX with embedded images.
        
        For DOCX files containing screenshots (not native tables),
        we extract images and process each separately.
        """
        from docx import Document
        
        logger.info(f"Extracting from DOCX: {docx_path}")
        
        doc = Document(docx_path)
        all_tables = []
        all_text = []
        total_confidence = 0.0
        image_count = 0
        
        # Extract embedded images
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_count += 1
                try:
                    image_data = rel.target_part.blob
                    image = Image.open(io.BytesIO(image_data))
                    
                    logger.info(f"Processing embedded image {image_count}")
                    result = self.extract_from_image(image)
                    
                    all_tables.extend(result.tables)
                    all_text.append(result.text)
                    total_confidence += result.confidence
                    
                except Exception as e:
                    logger.error(f"Failed to process image {image_count}: {e}")
        
        # Also extract any native text
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)
        
        avg_confidence = total_confidence / max(image_count, 1)
        
        return TableExtractionResult(
            tables=all_tables,
            text="\n".join(all_text),
            confidence=avg_confidence,
            metadata={
                "image_count": image_count,
                "source": "docx_embedded"
            }
        )
    
    @staticmethod
    def _raw_score_to_grade(score: float) -> str:
        """Map a raw 0-1 float score to a Docling-style grade string.

        Used as fallback when Docling's quality report grades are unavailable.
        """
        if score >= 0.85:
            return "EXCELLENT"
        elif score >= 0.65:
            return "GOOD"
        elif score >= 0.40:
            return "FAIR"
        else:
            return "POOR"

    def _process_result(self, docling_result) -> TableExtractionResult:
        """Process Docling result into our format with text post-processing."""
        tables = []

        # Import post-processor
        try:
            from preprocessing.text_postprocessor import postprocess_text_with_stats, postprocess_dataframe
            has_postprocessor = True
        except ImportError:
            has_postprocessor = False
            logger.warning("Text post-processor not available")

        # Track total corrections for QCS penalty
        total_corrections = 0
        correction_stats = {}

        # Extract tables as DataFrames
        doc = docling_result.document

        # Collect table confidence scores
        table_confidences = []

        for table in doc.tables:
            try:
                df = table.export_to_dataframe(doc=doc)

                # Apply text post-processing to fix missing spaces
                if has_postprocessor:
                    postprocess_dataframe(df, self.accounting_standard)

                tables.append(df)

                # Try to get table-level confidence from prov (provenance)
                if hasattr(table, 'prov') and table.prov:
                    for prov in table.prov:
                        if hasattr(prov, 'confidence') and prov.confidence is not None:
                            table_confidences.append(prov.confidence)

            except Exception as e:
                logger.warning(f"Failed to export table: {e}")

        # Extract text
        text = doc.export_to_markdown()

        # Post-process text and track corrections for QCS penalty
        if has_postprocessor:
            text, correction_stats = postprocess_text_with_stats(text, self.accounting_standard)
            total_corrections = correction_stats.get("total_corrections", 0)
            logger.info(f"Post-processing made {total_corrections} corrections")

        # =============================================================
        # Extract confidence scores and grades from Docling
        # =============================================================
        ocr_scores = []
        layout_scores = []

        # Check content items for scores
        for item in doc.texts:
            if hasattr(item, 'prov') and item.prov:
                for prov in item.prov:
                    if hasattr(prov, 'confidence') and prov.confidence is not None:
                        ocr_scores.append(prov.confidence)

        # Also check body for additional provenance
        if hasattr(doc, 'body') and doc.body:
            if hasattr(doc.body, 'children'):
                for child in doc.body.children:
                    if hasattr(child, 'prov') and child.prov:
                        for prov in child.prov:
                            if hasattr(prov, 'confidence') and prov.confidence is not None:
                                layout_scores.append(prov.confidence)

        # Calculate average scores
        avg_ocr = sum(ocr_scores) / len(ocr_scores) if ocr_scores else 0.0
        avg_layout = sum(layout_scores) / len(layout_scores) if layout_scores else 0.0

        # Combined confidence
        all_scores = ocr_scores + layout_scores + table_confidences
        avg_confidence = sum(all_scores) / len(all_scores) if all_scores else 0.0

        # =============================================================
        # Extract grades from Docling quality report (preferred)
        # Fallback: map raw average scores to grades
        # =============================================================
        ocr_grade = "FAIR"
        layout_grade = "FAIR"
        parse_grade = "FAIR"
        low_grade = "FAIR"

        # Try Docling's quality/confidence report (if available)
        quality_report = None
        for attr in ('quality', 'confidence_report', 'quality_report'):
            if hasattr(doc, attr):
                quality_report = getattr(doc, attr)
                break

        if quality_report is not None:
            # Extract grades from Docling's report
            for grade_attr, target in [
                ('ocr_grade', 'ocr'),
                ('layout_grade', 'layout'),
                ('parse_grade', 'parse'),
                ('low_grade', 'low'),
            ]:
                val = getattr(quality_report, grade_attr, None)
                if val is not None:
                    # Handle both string and enum values
                    grade_str = val.value if hasattr(val, 'value') else str(val).upper()
                    if grade_str in ("EXCELLENT", "GOOD", "FAIR", "POOR"):
                        if target == 'ocr':
                            ocr_grade = grade_str
                        elif target == 'layout':
                            layout_grade = grade_str
                        elif target == 'parse':
                            parse_grade = grade_str
                        elif target == 'low':
                            low_grade = grade_str
            logger.info(f"Docling grades: ocr={ocr_grade}, layout={layout_grade}, "
                        f"parse={parse_grade}, low={low_grade}")
        else:
            # Fallback: map raw scores to grades
            if avg_ocr > 0:
                ocr_grade = self._raw_score_to_grade(avg_ocr)
            if avg_layout > 0:
                layout_grade = self._raw_score_to_grade(avg_layout)
            # parse_grade stays FAIR (no raw parse score available from provenance)
            # low_grade: use the worst of ocr and layout grades
            grade_order = {"POOR": 0, "FAIR": 1, "GOOD": 2, "EXCELLENT": 3}
            low_grade = min(ocr_grade, layout_grade,
                           key=lambda g: grade_order.get(g, 1))
            logger.info(f"Fallback grades (from raw scores): ocr={ocr_grade}, "
                        f"layout={layout_grade}, low={low_grade}")

        logger.debug(f"Confidence scores - OCR: {avg_ocr:.3f}, Layout: {avg_layout:.3f}")

        return TableExtractionResult(
            tables=tables,
            text=text,
            confidence=avg_confidence,
            ocr_score=avg_ocr,
            layout_score=avg_layout,
            corrections_made=total_corrections,
            ocr_grade=ocr_grade,
            layout_grade=layout_grade,
            parse_grade=parse_grade,
            low_grade=low_grade,
            metadata={
                "table_count": len(tables),
                "postprocessed": has_postprocessor,
                "accounting_standard": self.accounting_standard,
                "ocr_samples": len(ocr_scores),
                "layout_samples": len(layout_scores),
                "table_samples": len(table_confidences),
                "correction_stats": correction_stats,
                "quality_report_available": quality_report is not None,
            }
        )


def extract_document(path: Path, language: str = "fr") -> TableExtractionResult:
    """
    Convenience function to extract from any supported document type.
    
    Args:
        path: Path to PDF, DOCX, or image file
        language: OCR language code
        
    Returns:
        TableExtractionResult
    """
    extractor = DoclingExtractor(language=language)
    
    suffix = path.suffix.lower()
    
    if suffix == '.pdf':
        return extractor.extract_from_pdf(path)
    elif suffix == '.docx':
        return extractor.extract_from_docx(path)
    elif suffix in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
        image = Image.open(path)
        return extractor.extract_from_image(image)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


if __name__ == "__main__":
    # Test extraction
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python docling_extractor.py <file_path>")
        sys.exit(1)
    
    path = Path(sys.argv[1])
    result = extract_document(path)
    
    print(f"\n{'='*60}")
    print(f"Extraction Result: {path.name}")
    print(f"{'='*60}")
    print(f"Tables found: {len(result.tables)}")
    print(f"Confidence: {result.confidence:.3f}")
    print(f"Text length: {len(result.text)} chars")
    
    for i, table in enumerate(result.tables):
        print(f"\nTable {i+1}:")
        print(table.head().to_string())
