"""
FinAlze OCR Pipeline - Main Orchestrator

Cascading OCR pipeline with QCS-based routing:
- Tier 0: Native extraction (PyMuPDF) — digital PDFs with embedded text
- Tier 1: Fast OCR (Docling + TableFormer)
- Tier 2: VLM fallback (Qwen2.5-VL) [placeholder]
- Tier 3: Ensemble / HITL
"""

import argparse
import logging
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def _all_canonical_table_objects(items: list) -> bool:
    """Return True iff every element is a live ``CanonicalTable`` (not a dict).

    Used to gate dual-tier reconciliation — older runs may have serialized
    Tier 0 ``canonical_tables`` to dicts, in which case reconciliation has
    nothing to merge against and we should fall through.
    """
    if not items:
        return False
    try:
        from accounting.canonical_model import CanonicalTable
    except ImportError:
        return False
    return all(isinstance(it, CanonicalTable) for it in items)


@dataclass
class ExtractionResult:
    """Result from OCR extraction.

    ``canonical_tables`` holds the progressively enriched output (§26 of the
    nomenclature revamp plan). Populated by Phase 2+; earlier tiers that
    haven't been migrated yet leave it as an empty list — legacy consumers
    of ``tables`` continue to work unchanged.
    """
    text: str
    tables: list
    qcs_score: float
    tier_used: int
    confidence_details: dict
    metadata: dict
    canonical_tables: list = field(default_factory=list)
    # Pre-reconciliation, per-engine canonical tables. Tier 3 ensemble voting
    # consumes these so it does not double-count Tier 0 (Tier 1's in-tier
    # dual-tier reconciliation overwrites ``canonical_tables``). Each tier sets
    # this to its own enrichment output before any cross-tier merge.
    raw_canonical_tables: list = field(default_factory=list)


@dataclass
class PipelineConfig:
    """Configuration for the OCR pipeline."""
    # QCS thresholds (grade-mapped intrinsic scores now use full 0-1 range)
    tau_native_high: float = 0.95
    tau_fast_high: float = 0.75  # Restored — grade mapping fixes threshold collapse
    tau_vlm_med: float = 0.50
    # Tier 3 ensemble routing on the Consensus-Entropy table quality Q_table
    # (arXiv 2504.11101): Q ≥ high → auto-accept; arbiter band → LLM-judge hook;
    # below arbiter (or any critical arithmetic failure) → HITL.
    tau_ensemble_high: float = 0.90
    tau_ensemble_arbiter: float = 0.70
    tau_ensemble_min: float = 0.60  # hard HITL floor (legacy)
    
    # Accounting standard
    accounting_standard: str = "IFRS"  # IFRS, NCT, SYSCOHADA
    
    # Performance
    enable_gpu: bool = True
    batch_size: int = 1
    
    # Output
    output_format: str = "json"  # json, markdown, csv
    # Document classification confidence threshold (0-1). If the
    # classifier returns a confidence below this value, the pipeline
    # must reject the document as non-financial per client spec.
    doc_type_confidence_threshold: float = 0.70
    # Score aggregation weights for Score 5 (s1..s4 must sum to 1.0)
    score_weights: dict = field(default_factory=lambda: {
        "s1": 0.25,  # Quality OCR
        "s2": 0.20,  # Identification/classification
        "s3": 0.25,  # Extraction confidence
        "s4": 0.30,  # Mapping/validation confidence
    })


class FinAlzePipeline:
    """Main OCR pipeline orchestrator."""
    
    def __init__(self, config: Optional[PipelineConfig] = None):
        self.config = config or PipelineConfig()
        self._initialize_models()
    
    def _initialize_models(self):
        """Lazy-load OCR models."""
        logger.info("Initializing FinAlze OCR Pipeline...")
        # Models will be loaded on first use
        self._docling_converter = None
        self._surya_predictor = None
        self._marker_converter = None
    
    def process_document(
        self,
        input_path: Path,
        preprocess_resize: bool = True,
        preprocess_grayscale: bool = True,
        max_image_width: int = 2000,
        accounting_standard: str = "NCT"
    ) -> ExtractionResult:
        """
        Process a single document through the cascading pipeline.
        
        Args:
            input_path: Path to PDF, DOCX, or image file
            preprocess_resize: Resize large images for faster processing
            preprocess_grayscale: Convert to grayscale for faster OCR
            max_image_width: Maximum image width before resizing
            accounting_standard: One of 'IFRS', 'NCT', 'SYSCOHADA'
            
        Returns:
            ExtractionResult with extracted text, tables, and QCS score
        """
        logger.info(f"Processing: {input_path}")
        
        # Track tier attempts for transparency
        tier_attempts = []
        
        # Step 1: Pre-processing
        # Strict classification (client rule): must be one of the four
        # financial categories with sufficient confidence, otherwise REJECT.
        classification = self._classify_document_strict(input_path)

        # Gate: reject if category is 'Autre' or confidence below threshold
        if (
            classification.get("category") == "Autre"
            or classification.get("confidence", 0.0)
            < self.config.doc_type_confidence_threshold
        ):
            logger.warning(
                "Document rejected by strict classifier: %s",
                classification,
            )
            # Persist a local audit event (synchronous) so the rejection is
            # immediately traceable without requiring async DB access.
            # Try async DB persistence first; if that fails, fall back to local JSONL
            try:
                import asyncio
                from pipeline_audit_db import async_log_rejection

                coro = async_log_rejection(
                    filename=str(input_path.name),
                    classification=classification,
                    threshold=self.config.doc_type_confidence_threshold,
                )
                try:
                    loop = asyncio.get_event_loop()
                    if loop.is_running():
                        loop.create_task(coro)
                    else:
                        asyncio.run(coro)
                except RuntimeError:
                    # No running loop — run synchronously
                    asyncio.run(coro)
            except Exception:
                try:
                    from pipeline_audit import log_rejection

                    log_rejection(str(input_path.name), classification, self.config.doc_type_confidence_threshold)
                except Exception:
                    logger.debug("Failed to persist audit event (db + jsonl) for rejection")

            # Log rejection in metadata and return a rejection ExtractionResult
            return ExtractionResult(
                text="",
                tables=[],
                qcs_score=0.0,
                tier_used=-1,
                confidence_details={"classification": classification},
                metadata={
                    "pipeline_status": "REJECTED",
                    "reject_reason": classification.get("reject_reason", ""),
                    "classification": classification,
                },
            )

        # Continue normally
        doc_type = classification.get("category")
        images = self._extract_images(input_path)

        logger.info(f"Document type: {doc_type}, Images: {len(images)}")
        
        # Step 2: Tier 0 - Native extraction (digital PDFs with embedded text)
        tier0_result = None
        if input_path.suffix.lower() == '.pdf':
            tier0_result = self._tier0_native(input_path)
            # §18 short-circuit: green-flag is the authoritative gate. QCS is
            # checked too as a guard against gross quality failures, but a
            # green-flag pass *plus* QCS ≥ tau_native_high is required to skip
            # Tier 1.
            green_passed = bool(tier0_result.metadata.get("green_flag_passed"))
            qcs_ok = tier0_result.qcs_score >= self.config.tau_native_high
            tier0_accepted = green_passed and qcs_ok

            tier_attempts.append({
                "tier": 0,
                "name": "Native PDF",
                "qcs": round(tier0_result.qcs_score, 3),
                "threshold": self.config.tau_native_high,
                "green_flag": green_passed,
                "passed": tier0_accepted,
            })
            if tier0_accepted:
                logger.info(
                    f"Tier 0 accepted (green_flag=PASS, QCS={tier0_result.qcs_score:.3f})"
                )
                tier0_result.confidence_details['tier_attempts'] = tier_attempts
                return tier0_result
            # Tier 0 below threshold — its findings will be forwarded to Tier 1
            if tier0_result.text or tier0_result.tables:
                gates_failed = tier0_result.metadata.get("green_flag_failures") or []
                logger.info(
                    f"Tier 0 below threshold (green_flag={'PASS' if green_passed else 'FAIL'}, "
                    f"QCS={tier0_result.qcs_score:.3f}) — forwarding to Tier 1 "
                    f"(text={len(tier0_result.text)} chars, tables={len(tier0_result.tables)}, "
                    f"gate_failures={len(gates_failed)})"
                )
        
        # Step 3: Tier 1 - Fast OCR (with Tier 0 context if available)
        result = self._tier1_fast_ocr(
            input_path, images, classification,
            preprocess_resize=preprocess_resize,
            preprocess_grayscale=preprocess_grayscale,
            max_image_width=max_image_width,
            accounting_standard=accounting_standard,
            tier0_result=tier0_result,
        )
        tier_attempts.append({
            "tier": 1, 
            "name": "Docling OCR", 
            "qcs": round(result.qcs_score, 3),
            "threshold": self.config.tau_fast_high,
            "passed": result.qcs_score >= self.config.tau_fast_high
        })
        if result.qcs_score >= self.config.tau_fast_high:
            logger.info(f"Tier 1 accepted (QCS={result.qcs_score:.3f})")
            result.confidence_details['tier_attempts'] = tier_attempts
            return result
        
        # Step 4: Tier 2 - VLM (placeholder for now)
        logger.info("Escalating to Tier 2 (VLM)...")
        tier1_result = result  # Keep Tier 1 result for the ensemble
        tier2_result = self._tier2_vlm(images)
        tier_attempts.append({
            "tier": 2,
            "name": "VLM",
            "qcs": round(tier2_result.qcs_score, 3),
            "threshold": self.config.tau_vlm_med,
            "passed": tier2_result.qcs_score >= self.config.tau_vlm_med
        })
        if tier2_result.qcs_score >= self.config.tau_vlm_med:
            logger.info(f"Tier 2 accepted (QCS={tier2_result.qcs_score:.3f})")
            tier2_result.confidence_details['tier_attempts'] = tier_attempts
            return tier2_result

        # Step 5: Tier 3 - Ensemble consensus (LV-ROVER vote + Consensus-Entropy routing)
        return self._tier3_ensemble(
            tier0_result=tier0_result,
            tier1_result=tier1_result,
            tier2_result=tier2_result,
            tier_attempts=tier_attempts,
        )
    
    def _classify_document(self, path: Path) -> str:
        """Classify a document's *dominant* statement type (§24).

        Samples text from up to the first 3 pages of a PDF (or the entire
        DOCX body), feeds rows to ``NomenclatureDictionary.classify_statement``.
        Returns ``"bilan"`` (backward-compatible) when classification is
        ambiguous — most provisional EF packs lead with the Bilan.
        """
        try:
            from accounting.nomenclature import load_default_dictionary
            dictionary = load_default_dictionary()
        except ImportError:
            return "bilan"

        rows: list[str] = []
        suffix = path.suffix.lower()

        if suffix == ".pdf":
            try:
                import fitz  # PyMuPDF
                with fitz.open(str(path)) as doc:
                    for page in list(doc)[:3]:
                        text = page.get_text() or ""
                        rows.extend(
                            line.strip() for line in text.splitlines() if line.strip()
                        )
            except Exception as exc:  # noqa: BLE001
                logger.debug(f"Document classification: PDF read failed — {exc}")
                return "bilan"
        elif suffix == ".docx":
            try:
                from docx import Document
                doc = Document(path)
                rows.extend(p.text.strip() for p in doc.paragraphs if p.text.strip())
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                rows.append(cell.text.strip())
            except Exception as exc:  # noqa: BLE001
                logger.debug(f"Document classification: DOCX read failed — {exc}")
                return "bilan"
        else:
            # Single-image inputs cannot be classified pre-OCR — default.
            return "bilan"

        classification = dictionary.classify_statement(rows)
        return classification if classification != "unknown" else "bilan"

    def _classify_document_strict(self, path: Path) -> dict:
        """Strict classifier wrapper returning JSON-like dict.

        Returns a dict: {"category": <one of the 4 values>,
        "confidence": <0-1>, "reject_reason": <str>}.
        Falls back to conservative outputs when the nomenclature
        classifier doesn't expose confidences.
        """
        # Default conservative response
        result = {"category": "Autre", "confidence": 0.0, "reject_reason": ""}
        try:
            from accounting.nomenclature import load_default_dictionary
            dictionary = load_default_dictionary()
        except Exception:
            result["reject_reason"] = "classifier_unavailable"
            return result

        # Reuse the lightweight label from the existing helper
        try:
            raw_label = self._classify_document(path)
        except Exception:
            raw_label = "unknown"

        low = (raw_label or "").lower()
        if "bilan" in low:
            mapped = "Bilan"
        elif "flux" in low:
            mapped = "Flux de trésorerie"
        elif "resultat" in low or "compte" in low:
            mapped = "Compte de résultat"
        else:
            mapped = "Autre"

        # Attempt to obtain a confidence score if the dictionary provides one
        confidence = 0.0
        try:
            # Collect a small sample of rows as in _classify_document
            rows: list[str] = []
            suffix = path.suffix.lower()
            if suffix == ".pdf":
                import fitz
                with fitz.open(str(path)) as doc:
                    for page in list(doc)[:3]:
                        text = page.get_text() or ""
                        rows.extend(line.strip() for line in text.splitlines() if line.strip())
            elif suffix == ".docx":
                from docx import Document
                doc = Document(path)
                rows.extend(p.text.strip() for p in doc.paragraphs if p.text.strip())
                for tbl in doc.tables:
                    for row in tbl.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                rows.append(cell.text.strip())

            # Various dictionary APIs may exist; try common patterns
            if hasattr(dictionary, "classify_statement_with_confidence"):
                lbl, confidence = dictionary.classify_statement_with_confidence(rows)
            elif hasattr(dictionary, "classify_statement_confidence"):
                res = dictionary.classify_statement_confidence(rows)
                if isinstance(res, dict):
                    lbl = res.get("label", raw_label)
                    confidence = res.get("confidence", 0.0)
                elif isinstance(res, tuple):
                    lbl, confidence = res
                else:
                    confidence = 0.0
            else:
                # No confidence API available: be conservative
                confidence = 0.5 if mapped != "Autre" else 0.0

        except Exception:
            confidence = 0.0

        result = {
            "category": mapped,
            "confidence": float(confidence),
            "reject_reason": "" if mapped != "Autre" else "not_a_financial_statement",
        }
        return result
    
    def _extract_images(self, path: Path) -> list:
        """Extract embedded images from DOCX or PDF pages."""
        images = []
        
        if path.suffix.lower() == '.docx':
            # Extract images from DOCX
            try:
                from docx import Document
                doc = Document(path)
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        images.append(rel.target_part.blob)
                logger.info(f"Extracted {len(images)} images from DOCX")
            except Exception as e:
                logger.error(f"DOCX image extraction failed: {e}")
        
        elif path.suffix.lower() in ['.jpg', '.jpeg', '.png']:
            # Direct image
            images.append(path.read_bytes())
        
        return images
    
    def _tier0_native(self, path: Path) -> ExtractionResult:
        """Tier 0: Native PDF text extraction using PyMuPDF.

        For digitally-generated PDFs (embedded text, not scanned images),
        this extracts text and tables natively — no OCR needed.

        Advantages over Tier 1 (Docling OCR):
        - Perfect text encoding (no garbled accents)
        - Native table extraction (no cell-reference issues)
        - Faster execution (no model inference)

        Falls through to Tier 1 if:
        - PDF is scanned (no embedded text)
        - Table extraction fails or finds no tables
        - QCS < tau_native_high (0.95)
        """
        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.warning("PyMuPDF not installed — skipping Tier 0")
            return ExtractionResult(
                text="", tables=[], qcs_score=0.0, tier_used=0,
                confidence_details={"status": "fitz_not_available"},
                metadata={}
            )

        try:
            pdf_doc = fitz.open(str(path))
            num_pages = len(pdf_doc)
            logger.info(f"Tier 0: Opened PDF with {num_pages} page(s)")

            # ---------------------------------------------------------------
            # Step 1: Detect digital vs scanned PDF
            # ---------------------------------------------------------------
            first_page_text = pdf_doc[0].get_text("text").strip()
            if len(first_page_text) < 50:
                logger.info("Tier 0: PDF appears scanned (little embedded text) — skipping")
                pdf_doc.close()
                return ExtractionResult(
                    text="", tables=[], qcs_score=0.0, tier_used=0,
                    confidence_details={"status": "scanned_pdf", "text_length": len(first_page_text)},
                    metadata={}
                )

            # ---------------------------------------------------------------
            # Step 2: Extract text from all pages
            # ---------------------------------------------------------------
            all_text_parts = []
            for page in pdf_doc:
                page_text = page.get_text("text")
                if page_text.strip():
                    all_text_parts.append(page_text)
            
            full_text = "\n\n".join(all_text_parts)
            logger.info(f"Tier 0: Extracted {len(full_text)} chars of native text")

            # ---------------------------------------------------------------
            # Step 3: Extract tables natively (two-strategy approach)
            # ---------------------------------------------------------------
            import pandas as pd
            native_tables = []

            def _table_looks_garbled(df: pd.DataFrame) -> bool:
                """Detect if find_tables() produced garbled output.

                Garbled tables have cells with multiple space-separated numbers
                that should be in separate columns (e.g. '95,276,113 41,694,461').
                """
                if len(df.columns) < 2:
                    return True  # single column when expecting multi-column

                import re
                multi_num_re = re.compile(
                    r'\d[\d\s.,]{3,}\s+\d[\d\s.,]{3,}'  # two number groups separated by space
                )
                garbled_cells = 0
                total_cells = 0

                for col_idx in range(1, len(df.columns)):  # skip label column
                    for val in df.iloc[:, col_idx]:
                        if pd.isna(val):
                            continue
                        val_str = str(val).strip()
                        if not val_str:
                            continue
                        total_cells += 1
                        if multi_num_re.search(val_str):
                            garbled_cells += 1

                # If >20% of non-empty cells have concatenated numbers, it's garbled
                return total_cells > 0 and garbled_cells / total_cells > 0.20

            for page_idx, page in enumerate(pdf_doc):
                try:
                    # Strategy 1: "lines" (default — uses vector graphics for grid)
                    tabs = page.find_tables()
                    page_tables = []
                    for tab_idx, tab in enumerate(tabs.tables):
                        df = tab.to_pandas()
                        if len(df) < 3:
                            continue
                        page_tables.append(df)

                    # Check if any table looks garbled
                    needs_retry = any(_table_looks_garbled(df) for df in page_tables)

                    if needs_retry and page_tables:
                        logger.info(
                            f"Tier 0: Page {page_idx+1} has garbled columns — "
                            f"retrying with strategy='text'"
                        )
                        # Strategy 2: "text" (uses text positions for column boundaries)
                        tabs_text = page.find_tables(strategy="text")
                        text_tables = []
                        for tab_idx, tab in enumerate(tabs_text.tables):
                            df = tab.to_pandas()
                            if len(df) < 3:
                                continue
                            text_tables.append(df)

                        # Use text-strategy tables if they have MORE columns (better split)
                        if text_tables:
                            lines_cols = max(len(df.columns) for df in page_tables)
                            text_cols = max(len(df.columns) for df in text_tables)
                            if text_cols > lines_cols:
                                page_tables = text_tables
                                logger.info(
                                    f"Tier 0: Text strategy improved: "
                                    f"{lines_cols} → {text_cols} columns"
                                )

                    for df in page_tables:
                        native_tables.append(df)
                        logger.info(
                            f"Tier 0: Page {page_idx+1}: "
                            f"{df.shape[0]} rows × {df.shape[1]} cols"
                        )
                except AttributeError:
                    logger.warning("Tier 0: find_tables() not available — upgrade PyMuPDF >= 1.23")
                    break
                except Exception as e:
                    logger.warning(f"Tier 0: Table extraction error on page {page_idx+1}: {e}")

            pdf_doc.close()

            if not native_tables:
                logger.info("Tier 0: No tables found — falling through to Tier 1")
                return ExtractionResult(
                    text=full_text, tables=[], qcs_score=0.0, tier_used=0,
                    confidence_details={"status": "no_tables_found", "text_length": len(full_text)},
                    metadata={"native_text_available": True}
                )

            logger.info(f"Tier 0: Found {len(native_tables)} table(s) natively")

            # ---------------------------------------------------------------
            # Step 3b: Explode merged cells
            # ---------------------------------------------------------------
            # Some accounting PDFs have physically merged table cells where
            # the entire body is in one row with \n-separated values.
            # Detect and explode these into proper row-per-line-item DFs.

            def _explode_merged_rows(df: pd.DataFrame) -> pd.DataFrame:
                """Explode a DataFrame whose body is collapsed into merged cells.

                Pattern: row 0 = header, row 1 = mega-cell with \\n-separated
                values across all columns, row 2+ = totals. The mega-cell is
                detected when any cell has >= 3 newline characters.
                """
                new_rows = []
                for _, row in df.iterrows():
                    # Check if any cell in this row has multiple newlines
                    max_parts = 1
                    for val in row:
                        if pd.notna(val) and isinstance(val, str) and '\n' in val:
                            parts_count = len(val.split('\n'))
                            max_parts = max(max_parts, parts_count)

                    if max_parts <= 2:
                        # Normal row — pass through
                        new_rows.append(row.tolist())
                    else:
                        # Merged row — explode into max_parts rows
                        split_cols = []
                        for val in row:
                            if pd.notna(val) and isinstance(val, str) and '\n' in val:
                                split_cols.append(val.split('\n'))
                            else:
                                # Single value — pad with empty strings
                                split_cols.append([str(val) if pd.notna(val) else ''] * max_parts)

                        # Pad shorter columns to max_parts
                        for i, parts in enumerate(split_cols):
                            if len(parts) < max_parts:
                                split_cols[i] = parts + [''] * (max_parts - len(parts))

                        for row_idx in range(max_parts):
                            new_rows.append([col[row_idx] for col in split_cols])

                result = pd.DataFrame(new_rows, columns=df.columns)
                return result

            exploded_tables = []
            for i, df in enumerate(native_tables):
                # Detect: any cell with >= 3 newlines indicates a merged table
                has_merged = False
                for col in df.columns:
                    for val in df[col]:
                        if pd.notna(val) and isinstance(val, str) and val.count('\n') >= 3:
                            has_merged = True
                            break
                    if has_merged:
                        break

                if has_merged:
                    exploded = _explode_merged_rows(df)
                    logger.info(
                        f"Tier 0: Table {i+1} had merged cells — "
                        f"exploded {df.shape[0]} rows → {exploded.shape[0]} rows"
                    )
                    exploded_tables.append(exploded)
                else:
                    exploded_tables.append(df)

            native_tables = exploded_tables

            # ---------------------------------------------------------------
            # Step 4: Post-processing (same as Tier 1)
            # ---------------------------------------------------------------
            total_corrections = 0
            correction_stats = {}
            standard = self.config.accounting_standard

            try:
                from preprocessing.text_postprocessor import postprocess_text_with_stats, postprocess_dataframe
                full_text, correction_stats = postprocess_text_with_stats(full_text, standard)
                total_corrections = correction_stats.get("total_corrections", 0)
                for df in native_tables:
                    postprocess_dataframe(df, standard)
                logger.info(f"Tier 0: Post-processing made {total_corrections} corrections")
            except ImportError:
                logger.warning("Tier 0: Text post-processor not available")

            # ---------------------------------------------------------------
            # Step 4b: Enrichment Pipeline (§23 DAG Stages 1–10)
            # ---------------------------------------------------------------
            canonical_tables = []
            green_flag_result = None
            try:
                from accounting.enrichment_pipeline import enrich_tables
                canonical_tables = enrich_tables(
                    raw_tables=native_tables,
                    tier=0,
                )
                logger.info(
                    f"Tier 0: Enrichment pipeline produced "
                    f"{len(canonical_tables)} canonical tables"
                )

                # Green-flag gate (§18) — check if Tier 0 quality is sufficient
                if canonical_tables:
                    try:
                        from accounting.green_flag import check_green_flag
                        # Compute metrics from enrichment results
                        total_rows = sum(len(ct.rows) for ct in canonical_tables)
                        matched_rows = sum(
                            sum(1 for r in ct.rows
                                if r.match_type in ("exact", "fuzzy", "resolved"))
                            for ct in canonical_tables
                        )
                        match_rate = matched_rows / max(total_rows, 1)

                        total_numeric = sum(
                            sum(1 for c in r.cells.values() if c.parsed_value is not None)
                            for ct in canonical_tables for r in ct.rows
                        )
                        parse_fails = sum(
                            sum(1 for c in r.cells.values()
                                if c.flag == "yellow" and c.raw_value.strip())
                            for ct in canonical_tables for r in ct.rows
                        )
                        parse_fail_rate = parse_fails / max(total_numeric, 1)

                        col_conf = max(
                            (ct.metadata.get("column_confidence", 0.0)
                             for ct in canonical_tables),
                            default=0.0,
                        ) if any(ct.metadata for ct in canonical_tables) else 0.80

                        class_conf = max(
                            (ct.match_rate for ct in canonical_tables),
                            default=0.0,
                        )

                        green_flag_result = check_green_flag(
                            match_rate=match_rate,
                            parse_fail_rate=parse_fail_rate,
                            column_confidence=col_conf,
                            classification_confidence=class_conf,
                        )
                        logger.info(
                            f"Tier 0 green-flag: {'PASSED' if green_flag_result.passed else 'FAILED'} "
                            f"(match={match_rate:.0%}, parse_fail={parse_fail_rate:.0%})"
                        )
                    except ImportError:
                        logger.debug("Tier 0: Green-flag module not available")
            except ImportError:
                logger.debug("Tier 0: Enrichment pipeline not available")

            # ---------------------------------------------------------------
            # Step 5: Run V1-V4 validation (same as Tier 1)
            # ---------------------------------------------------------------
            validation_report = None
            try:
                from accounting.rule_engine import validate_with_rules
                validation_report = validate_with_rules(
                    tables=native_tables,
                    standard=standard,
                    run_v2=True,
                    run_v3=True,
                    run_v4=True,
                )
                logger.info(
                    f"Tier 0 validation: {validation_report.passed_checks}/"
                    f"{validation_report.total_checks} passed, "
                    f"critical={validation_report.critical_failures}"
                )
            except ImportError:
                try:
                    from accounting.validator import validate_extraction
                    validation_report = validate_extraction(native_tables)
                except ImportError:
                    logger.warning("Tier 0: No validator available")

            # ---------------------------------------------------------------
            # Step 6: Calculate QCS with EXCELLENT intrinsic grades
            # (native text = perfect OCR, layout, and parse quality)
            # ---------------------------------------------------------------
            try:
                from qcs.calculator import calculate_qcs

                qcs_report = calculate_qcs(
                    ocr_grade="EXCELLENT",
                    layout_grade="EXCELLENT",
                    parse_grade="EXCELLENT",
                    low_grade="EXCELLENT",
                    text=full_text,
                    tables=native_tables,
                    validation_report=validation_report,
                    corrections_made=total_corrections,
                )
                qcs_score = qcs_report.qcs_score
                qcs_grade = qcs_report.grade
                qcs_details = qcs_report.to_dict()
            except ImportError:
                qcs_score = 0.95 if validation_report and validation_report.overall_passed else 0.5
                qcs_grade = "excellent" if qcs_score >= 0.90 else "good"
                qcs_details = {}

            # Convert tables to list of dicts (same format as Tier 1)
            tables_data = [df.to_dict(orient='records') for df in native_tables]

            # Determine HITL from QCS report or validation
            hitl_required = False
            hitl_reason = ""
            if hasattr(qcs_report, 'needs_hitl'):
                hitl_required = qcs_report.needs_hitl
                hitl_reason = qcs_report.hitl_reason

            return ExtractionResult(
                text=full_text,
                tables=tables_data,
                qcs_score=qcs_score,
                tier_used=0,
                confidence_details={
                    'extraction_method': 'native_pymupdf',
                    'ocr_grade': 'EXCELLENT',
                    'layout_grade': 'EXCELLENT',
                    'parse_grade': 'EXCELLENT',
                    'low_grade': 'EXCELLENT',
                    'qcs_grade': qcs_grade,
                    'hitl_reason': hitl_reason,
                    **qcs_details,
                },
                metadata={
                    'doc_type': 'native_pdf',
                    'table_count': len(native_tables),
                    'page_count': num_pages,
                    'text_length': len(full_text),
                    'native_extraction': True,
                    'hitl_required': hitl_required,
                    'accounting_standard': standard,
                    'validation_report': validation_report.to_dict() if validation_report else None,
                    'correction_stats': correction_stats,
                    'green_flag_passed': green_flag_result.passed if green_flag_result else None,
                    'green_flag_failures': green_flag_result.failure_reasons if green_flag_result else [],
                },
                canonical_tables=canonical_tables,
                raw_canonical_tables=canonical_tables,
            )

        except Exception as e:
            import traceback
            logger.error(f"Tier 0 failed: {e}\n{traceback.format_exc()}")
            return ExtractionResult(
                text="", tables=[], qcs_score=0.0, tier_used=0,
                confidence_details={"error": str(e)},
                metadata={"failed": True}
            )
    
    def _tier1_fast_ocr(
        self,
        path: Path,
        images: list,
        doc_classification: dict,
        preprocess_resize: bool = True,
        preprocess_grayscale: bool = True,
        max_image_width: int = 2000,
        accounting_standard: str = "NCT",
        tier0_result: Optional[ExtractionResult] = None,
    ) -> ExtractionResult:
        """Tier 1: Fast OCR with Docling/TableFormer.

        If tier0_result is provided (from a sub-threshold Tier 0 run), its
        findings are leveraged:
        - Native text feeds the QCS heuristic layer (clean encoding)
        - Tables missing from Docling are supplemented from Tier 0
        """
        try:
            from extraction.docling_extractor import DoclingExtractor
            
            logger.info(f"Tier 1: Processing with Docling (standard={accounting_standard})...")
            extractor = DoclingExtractor(
                language="fr",
                preprocess_resize=preprocess_resize,
                preprocess_grayscale=preprocess_grayscale,
                max_image_width=max_image_width,
                accounting_standard=accounting_standard
            )
            
            # Choose extraction method based on file type
            suffix = path.suffix.lower()
            if suffix == '.pdf':
                result = extractor.extract_from_pdf(path)
            elif suffix == '.docx':
                result = extractor.extract_from_docx(path)
            elif suffix in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                from PIL import Image
                image = Image.open(path)
                result = extractor.extract_from_image(image)
            else:
                raise ValueError(f"Unsupported file type: {suffix}")
            
            # =============================================================
            # Leverage Tier 0 findings if available
            # =============================================================
            tier0_supplemented = 0
            tier0_text_used = False

            if tier0_result and tier0_result.tables and len(tier0_result.tables) > len(result.tables):
                # Tier 0 found tables that Docling missed — supplement
                import pandas as pd
                tier0_tables_as_df = []
                for t0_table in tier0_result.tables:
                    if isinstance(t0_table, dict):
                        # Already dict-of-records from Tier 0 serialization
                        tier0_tables_as_df.append(pd.DataFrame(t0_table))
                    elif isinstance(t0_table, list):
                        tier0_tables_as_df.append(pd.DataFrame(t0_table))
                    elif isinstance(t0_table, pd.DataFrame):
                        tier0_tables_as_df.append(t0_table)

                if len(tier0_tables_as_df) > len(result.tables):
                    # Append Tier 0 tables that Docling didn't find
                    # (Tier 0 extracted more tables natively)
                    for extra_df in tier0_tables_as_df[len(result.tables):]:
                        result.tables.append(extra_df)
                        tier0_supplemented += 1
                    logger.info(
                        f"Tier 1: Supplemented {tier0_supplemented} table(s) from Tier 0 native extraction"
                    )

            # Convert tables to list of dicts for JSON serialization
            tables_data = []
            for df in result.tables:
                tables_data.append(df.to_dict(orient='records'))

            # =============================================================
            # Enrichment Pipeline (§23 DAG Stages 1–10) — Tier 1
            # =============================================================
            canonical_tables_t1: list = []
            try:
                from accounting.enrichment_pipeline import enrich_tables
                if result.tables:
                    canonical_tables_t1 = enrich_tables(
                        raw_tables=result.tables,
                        tier=1,
                    )
                    logger.info(
                        f"Tier 1: Enrichment pipeline produced "
                        f"{len(canonical_tables_t1)} canonical tables"
                    )
            except ImportError:
                logger.debug("Tier 1: Enrichment pipeline not available")
            except Exception as exc:  # noqa: BLE001 — keep tier alive
                logger.warning(f"Tier 1: Enrichment failed — {exc}")

            # Preserve Tier 1's own enrichment output *before* any cross-tier
            # merge — the Tier 3 ensemble votes over genuine per-engine tables,
            # not the already-reconciled product (which would double-count Tier 0).
            raw_canonical_tables_t1 = list(canonical_tables_t1)

            # =============================================================
            # Dual-tier reconciliation (§10 + §20) — Tier 0 ⨝ Tier 1
            # When Tier 0 produced canonical_tables (digital PDF), merge
            # them with Tier 1's results: trust Tier 0 numbers, flag
            # disagreements, attach a per-cell conflict report.
            # =============================================================
            tier0_canonical = []
            if tier0_result is not None:
                tier0_canonical = list(tier0_result.canonical_tables or [])
                # Skip reconciliation if Tier 0 yielded only serialized dicts
                # (e.g. older runs) — we need live CanonicalTable objects.
                if tier0_canonical and not _all_canonical_table_objects(tier0_canonical):
                    logger.debug(
                        "Tier 1: Tier 0 canonical_tables are serialized — "
                        "skipping reconciliation"
                    )
                    tier0_canonical = []

            reconciliation_report: dict = {}
            if tier0_canonical and canonical_tables_t1:
                try:
                    from accounting.reconciliation import reconcile_dual_tier
                    reconciled = reconcile_dual_tier(
                        tier0_tables=tier0_canonical,
                        tier1_tables=canonical_tables_t1,
                    )
                    canonical_tables_t1 = reconciled
                    total_conflicts = sum(
                        len(ct.conflicts or []) for ct in reconciled
                    )
                    avg_coverage = (
                        sum(ct.metadata.get("alignment_coverage", 0.0) for ct in reconciled)
                        / max(len(reconciled), 1)
                    )
                    reconciliation_report = {
                        "applied": True,
                        "table_count": len(reconciled),
                        "conflict_count": total_conflicts,
                        "avg_alignment_coverage": round(avg_coverage, 3),
                    }
                    logger.info(
                        f"Tier 1: Dual-tier reconciliation produced "
                        f"{len(reconciled)} merged tables, "
                        f"{total_conflicts} cell conflicts, "
                        f"avg_coverage={avg_coverage:.0%}"
                    )
                except ImportError:
                    logger.debug("Tier 1: Reconciliation module not available")
                except Exception as exc:  # noqa: BLE001
                    logger.warning(f"Tier 1: Reconciliation failed — {exc}")

            # =============================================================
            # Run full validation hierarchy (V1-V4) with statement segmentation
            # Uses configurable rule engine with JSON rule definitions
            # =============================================================
            validation_report = None
            try:
                from accounting.rule_engine import validate_with_rules
                if result.tables:
                    validation_report = validate_with_rules(
                        tables=result.tables,
                        standard=accounting_standard,
                        run_v2=True,   # Cross-statement coherence
                        run_v3=True,   # Anomaly detection
                        run_v4=True,   # Ratio computation
                    )
                    logger.info(
                        f"Validation (V1-V4): {validation_report.passed_checks}/{validation_report.total_checks} passed, "
                        f"critical_failures={validation_report.critical_failures}, "
                        f"severity_weighted={validation_report.severity_weighted_score}"
                    )
            except ImportError:
                # Fallback to basic validator if rule engine not available
                try:
                    from accounting.validator import validate_extraction
                    if result.tables:
                        validation_report = validate_extraction(result.tables)
                except ImportError:
                    logger.warning("No validator available — semantic layer will default to 0.5")

            # =============================================================
            # Calculate QCS using grade-mapped 3-layer architecture
            # Includes: low_grade pre-gate, severity-weighted semantic,
            #           critical_failures HITL override
            # =============================================================
            try:
                from qcs.calculator import calculate_qcs

                qcs_report = calculate_qcs(
                    # Layer 1: Intrinsic (grade-mapped, table_score excluded)
                    ocr_grade=result.ocr_grade,
                    layout_grade=result.layout_grade,
                    parse_grade=result.parse_grade,
                    # Pre-gate
                    low_grade=result.low_grade,
                    # Layer 2: Heuristic (from OCR text — better accent handling)
                    text=result.text,
                    tables=result.tables,
                    # Layer 3: Semantic (severity-weighted validation)
                    validation_report=validation_report,
                    # Correction penalty
                    corrections_made=result.corrections_made,
                )

                qcs_score = qcs_report.qcs_score
                qcs_grade = qcs_report.grade
                qcs_details = qcs_report.to_dict()

            except ImportError:
                # Fallback if QCS module not available
                qcs_score = result.confidence if result.confidence > 0 else 0.5
                qcs_grade = "unknown"
                qcs_details = {}

                # Build scores (placeholders if detailed signals absent)
                doc_type_label = doc_classification.get("category") if isinstance(doc_classification, dict) else None

                # if result lacks confidence attributes, fallback gracefully
                docling_conf = getattr(result, "confidence", 0.5)
                try:
                    val_score = validation_report.severity_weighted_score
                except Exception:
                    val_score = 0.5 if validation_report else 0.5

                # Compute normalized float scores (0..1)
                s1 = float(qcs_score) if qcs_score is not None else 0.0
                s2 = float(doc_classification.get("confidence", 0.0)) if isinstance(doc_classification, dict) else 0.0
                s3 = float(docling_conf) if docling_conf is not None else 0.0
                s4 = float(val_score) if val_score is not None else 0.0

                weights = getattr(self.config, "score_weights", {"s1": 0.25, "s2": 0.2, "s3": 0.25, "s4": 0.3})
                # Ensure weights sum to 1.0 (normalize if not)
                total_w = sum(weights.values()) if isinstance(weights, dict) and weights else 1.0
                if total_w <= 0:
                    total_w = 1.0

                w1 = weights.get("s1", 0.25) / total_w
                w2 = weights.get("s2", 0.20) / total_w
                w3 = weights.get("s3", 0.25) / total_w
                w4 = weights.get("s4", 0.30) / total_w

                score5 = (w1 * s1 + w2 * s2 + w3 * s3 + w4 * s4) * 100.0

                scores = {
                    "score1": int(round(s1 * 100)),
                    "score2": int(round(s2 * 100)),
                    "score3": int(round(s3 * 100)),
                    "score4": int(round(s4 * 100)),
                    "score5": int(round(score5)),
                }

                # Synthetic bounding boxes if extractor doesn't provide them
                bboxes = result.metadata.get("bboxes") if isinstance(result.metadata, dict) and result.metadata.get("bboxes") else []
                if not bboxes and result.tables:
                    # create simple normalized boxes per table for UI overlay
                    bboxes = []
                    for i in range(len(result.tables)):
                        top = 0.05 + 0.08 * i
                        bottom = min(0.95, top + 0.12)
                        bboxes.append({
                            "page": 1,
                            "x0": 0.05,
                            "y0": round(top, 3),
                            "x1": 0.95,
                            "y1": round(bottom, 3),
                            "label": f"table_{i+1}",
                        })

                return ExtractionResult(
                text=result.text,
                tables=tables_data,
                qcs_score=qcs_score,
                tier_used=1,
                confidence_details={
                    'docling_confidence': result.confidence,
                    'ocr_grade': result.ocr_grade,
                    'layout_grade': result.layout_grade,
                    'parse_grade': result.parse_grade,
                    'low_grade': result.low_grade,
                    'qcs_grade': qcs_grade,
                    'scores': scores,
                    **qcs_details,
                },
                metadata={
                    'doc_type': doc_type_label,
                    'table_count': len(result.tables),
                    'preprocess_resize': preprocess_resize,
                    'preprocess_grayscale': preprocess_grayscale,
                    'validation_report': validation_report.to_dict() if validation_report else None,
                    'tier0_text_used': tier0_text_used,
                    'tier0_tables_supplemented': tier0_supplemented,
                    'reconciliation': reconciliation_report,
                    'bboxes': bboxes,
                    **result.metadata
                },
                canonical_tables=canonical_tables_t1,
                raw_canonical_tables=raw_canonical_tables_t1,
            )
        except Exception as e:
            import traceback
            logger.error(f"Tier 1 extraction failed: {e}\n{traceback.format_exc()}")
            # Return low-confidence result to trigger escalation
            return ExtractionResult(
                text="",
                tables=[],
                qcs_score=0.0,
                tier_used=1,
                confidence_details={'error': str(e)},
                metadata={'doc_type': (doc_classification.get('category') if isinstance(doc_classification, dict) else None), 'failed': True}
            )
    
    def _tier2_vlm(self, images: list) -> ExtractionResult:
        """Tier 2: VLM-based OCR."""
        # NOT IMPLEMENTED — return QCS=0 so it never falsely passes
        logger.warning("Tier 2 VLM is not implemented — returning QCS=0")
        return ExtractionResult(
            text="",
            tables=[],
            qcs_score=0.0,
            tier_used=2,
            confidence_details={"status": "not_implemented"},
            metadata={"vlm_implemented": False}
        )

    def _tier3_ensemble(
        self,
        tier0_result: Optional[ExtractionResult],
        tier1_result: ExtractionResult,
        tier2_result: Optional[ExtractionResult],
        tier_attempts: list,
    ) -> ExtractionResult:
        """Tier 3: N-way ensemble consensus (LV-ROVER vote + Consensus-Entropy routing).

        Votes the per-engine canonical tables (Tier 0 native, Tier 1 OCR, Tier 2
        VLM) into a single merged set, scores cross-engine agreement with the
        Consensus-Entropy table quality ``Q_table`` (arXiv 2504.11101), and routes:

            * < 2 engine sources      → HITL (no corroboration earns a lower bar)
            * critical arithmetic fail → HITL (overrides Q_table; from V1-V4 rules)
            * Q_table ≥ tau_ensemble_high     → auto-accept
            * tau_ensemble_arbiter ≤ Q < high → LLM-as-Judge band (no judge yet → HITL)
            * Q_table < tau_ensemble_arbiter  → HITL

        Replaces the legacy "pick best prior result + always-HITL" stub.
        """
        from accounting.ensemble import (
            EnsembleSource, ensemble_vote, canonical_tables_to_records,
        )

        # Assemble engine sources from each tier's *pre-reconciliation* canonical
        # tables (live CanonicalTable objects only — see _all_canonical_table_objects).
        weighted = [("tier0", 1.0, tier0_result), ("tier1", 0.8, tier1_result),
                    ("tier2", 0.7, tier2_result)]
        sources = []
        for name, weight, res in weighted:
            if res is None:
                continue
            raw = list(res.raw_canonical_tables or [])
            if raw and _all_canonical_table_objects(raw):
                sources.append(EnsembleSource(name=name, tables=raw, weight=weight))

        priors = [r for r in (tier0_result, tier1_result, tier2_result) if r is not None]
        best_prior = max(priors, key=lambda r: r.qcs_score) if priors else tier1_result

        # A critical arithmetic failure from any prior tier forces HITL regardless
        # of consensus (a balanced-looking-but-wrong table is the worst outcome).
        arithmetic_fail = any(
            ((r.metadata or {}).get("validation_report") or {}).get("critical_failures")
            for r in priors
        )

        source_count = len(sources)

        # ---- Degenerate: fewer than 2 engines → no consensus possible --------
        if source_count < 2:
            logger.warning(
                "Tier 3: %d engine source(s) — no consensus possible, routing to HITL",
                source_count,
            )
            result = best_prior
            tier_attempts.append({
                "tier": 3, "name": "Ensemble",
                "qcs": round(result.qcs_score, 3),
                "threshold": self.config.tau_ensemble_arbiter,
                "passed": False,
                "sources": [s.name for s in sources],
                "single_source": True,
            })
            result.metadata = {
                **result.metadata, "hitl_required": True,
                "ensemble": {"applied": False, "reason": "single_source",
                             "source_count": source_count},
            }
            result.confidence_details['tier_attempts'] = tier_attempts
            return result

        # ---- ≥2 engines → vote ----------------------------------------------
        voted = ensemble_vote(sources)
        # Conservative routing on the *worst* voted table's Q_table.
        q_tables = [
            ct.metadata.get("q_table", 1.0) for ct in voted
            if ct.metadata.get("reconciliation") == "ensemble"
        ]
        q_table = min(q_tables) if q_tables else best_prior.qcs_score
        total_conflicts = sum(len(ct.conflicts or []) for ct in voted)

        if arithmetic_fail or q_table < self.config.tau_ensemble_arbiter:
            passed, reason = False, ("arithmetic_failure" if arithmetic_fail else "low_consensus")
        elif q_table >= self.config.tau_ensemble_high:
            passed, reason = True, "high_consensus"
        else:
            passed, reason = False, "arbiter_band_no_judge"  # LLM judge not wired yet

        tier_attempts.append({
            "tier": 3, "name": "Ensemble",
            "qcs": round(q_table, 3),
            "threshold": self.config.tau_ensemble_high,
            "passed": passed,
            "sources": [s.name for s in sources],
            "conflicts": total_conflicts,
            "q_table": round(q_table, 3),
        })

        logger.info(
            "Tier 3 ensemble: %d sources, Q_table=%.3f, conflicts=%d → %s (%s)",
            source_count, q_table, total_conflicts,
            "ACCEPT" if passed else "HITL", reason,
        )

        return ExtractionResult(
            text=best_prior.text,
            tables=canonical_tables_to_records(voted),
            qcs_score=q_table,
            tier_used=3,
            confidence_details={
                **best_prior.confidence_details,
                "ensemble_q_table": round(q_table, 4),
                "ensemble_sources": [s.name for s in sources],
                "tier_attempts": tier_attempts,
            },
            metadata={
                **best_prior.metadata,
                "tier_used": 3,
                "hitl_required": not passed,
                "ensemble": {
                    "applied": True,
                    "source_count": source_count,
                    "source_names": [s.name for s in sources],
                    "q_table": round(q_table, 4),
                    "conflict_count": total_conflicts,
                    "routing_reason": reason,
                    "arithmetic_override": arithmetic_fail,
                },
            },
            canonical_tables=voted,
            raw_canonical_tables=best_prior.raw_canonical_tables,
        )


def main():
    parser = argparse.ArgumentParser(description='FinAlze OCR Pipeline')
    parser.add_argument('--input', '-i', required=True, help='Input file path')
    parser.add_argument('--output', '-o', default='output/', help='Output directory')
    parser.add_argument('--standard', '-s', default='IFRS',
                        choices=['IFRS', 'NCT', 'SYSCOHADA'],
                        help='Accounting standard')
    parser.add_argument('--remote', metavar='API_URL', default=None,
                        help='Submit to a FinAlze API (e.g. http://localhost:8000) instead of running locally')
    parser.add_argument('--api-key', default=None, help='Partner API key for --remote')
    parser.add_argument('--token', default=None, help='Bearer token for --remote')
    args = parser.parse_args()

    # Remote mode: submit to the backend API and poll for the result.
    if args.remote:
        import sys as _sys
        _sys.path.insert(0, str(Path(__file__).resolve().parent))
        from client.sdk import FinAlzeClient
        with FinAlzeClient(base_url=args.remote, api_key=args.api_key, token=args.token) as client:
            result_dict = client.submit_and_wait(args.input, accounting_standard=args.standard)
        print(f"\n{'='*60}\nFinAlze OCR Result (remote)\n{'='*60}")
        print(f"Input: {Path(args.input).name}")
        print(f"Tier Used: {result_dict.get('tier_used')}")
        print(f"QCS Score: {result_dict.get('qcs_score')}")
        print(f"HITL Required: {(result_dict.get('metadata') or {}).get('hitl_required', False)}")
        print(f"{'='*60}")
        return

    # Initialize pipeline
    config = PipelineConfig(accounting_standard=args.standard)
    pipeline = FinAlzePipeline(config)

    # Process document
    input_path = Path(args.input)
    result = pipeline.process_document(input_path)
    
    # Output
    print(f"\n{'='*60}")
    print(f"FinAlze OCR Result")
    print(f"{'='*60}")
    print(f"Input: {input_path.name}")
    print(f"Tier Used: {result.tier_used}")
    print(f"QCS Score: {result.qcs_score:.3f}")
    print(f"HITL Required: {result.metadata.get('hitl_required', False)}")
    print(f"{'='*60}")


if __name__ == "__main__":
    main()
